import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Define Paths
DOCX_PATH = r"c:\Users\ASUS\OneDrive\Desktop\someapp\CivicFlow_User_Guide.docx"
PDF_PATH = r"c:\Users\ASUS\OneDrive\Desktop\someapp\CivicFlow_User_Guide.pdf"


def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def create_docx_guide():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("CivicFlow Platform")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Navy/Indigo

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle_p.add_run("Official User Guide & Operational Manual")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Overview
    h1 = doc.add_heading("1. Executive Overview", level=1)
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    p = doc.add_paragraph(
        "CivicFlow is a state-of-the-art, citizen-driven civic accountability and infrastructure management platform. "
        "Unlike traditional portals where complaints disappear into black-hole queues, CivicFlow enforces complete transparency, "
        "AI-assisted classification, real-time SLA countdown timers, multi-level automatic supervisor escalations, "
        "and a mandatory citizen verification loop."
    )
    p.runs[0].font.size = Pt(11)

    # Key Features Bullet List
    features = [
        ("AI Auto-Classification", "Uses Gemini AI & rule-based engine to analyze complaint text/photos, assign severity, estimate SLA, and dispatch to the correct department."),
        ("SLA Countdown & Auto-Escalation", "Every issue is tracked with exact deadlines. Overdue issues automatically escalate to Department Supervisors and Municipal Commissioners."),
        ("Proof-of-Resolution", "Department officers must upload before/after photos and notes before marking an issue resolved."),
        ("Citizen Verification Loop", "The original reporter inspects the work and either closes the ticket or reopens it with immediate escalation."),
        ("Executive Command Center", "Provides platform administrators with real-time KPI metrics, volume distribution charts, and AI predictive insights.")
    ]
    for title, desc in features:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{title}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(15, 23, 42)
        r2 = bp.add_run(desc)
        r2.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Demo Credentials Table
    h2 = doc.add_heading("2. Access Roles & Demo Credentials", level=1)
    h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Role Type", "Demo Email", "Password", "Primary Responsibilities"]
    hdr_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        hdr_cells[idx].text = header_text
        set_cell_background(hdr_cells[idx], "1E3A8A")
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    role_data = [
        ("Citizen Reporter", "citizen@civicflow.org", "password", "Report issues, upvote community items, verify fixes, earn badges."),
        ("Department Officer", "authority@civicflow.org", "password", "Accept assigned tickets, dispatch crews, upload resolution photos."),
        ("Platform Administrator", "admin@civicflow.org", "password", "Monitor command center KPIs, review SLA breach alerts & escalations.")
    ]

    for row_idx, data in enumerate(role_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F8FAFC")
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Step-by-Step User Workflows
    h3 = doc.add_heading("3. Step-by-Step User Workflows", level=1)
    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    # Workflow 1
    doc.add_heading("3.1 Citizen Reporter Workflow", level=2)
    c_steps = [
        "Sign In: Click 'Sign In' at the top right and select 'Citizen User' quick login.",
        "File Report: Click 'Report Issue' in navigation to open the 4-step wizard.",
        "Step 1 - Photo Evidence: Drag and drop a photo, paste an image URL, or click any Quick Sample Photo preset (Pothole, Garbage, Streetlight, etc.).",
        "Step 2 - Pinpoint Location: Select exact coordinates on the interactive Leaflet GIS map.",
        "Step 3 - Description & AI: Describe the problem and run AI classification.",
        "Step 4 - AI Review & Duplicate Check: Review AI category/severity ratings. If duplicates exist nearby, click 'Support' to upvote existing report.",
        "Step 5 - Submit: Submit complaint to generate unique tracking code (e.g. CF-20260815-100001) and earn +15 gamification points."
    ]
    for s in c_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(s).font.size = Pt(10)

    # Workflow 2
    doc.add_heading("3.2 Department Officer Workflow", level=2)
    o_steps = [
        "Sign In: Log in as 'Officer User' (authority@civicflow.org).",
        "Dashboard Access: Open 'Dashboard' to view incoming department assignments.",
        "Accept Ticket: Select an 'Assigned' complaint and click 'Accept Assignment' to transition status to 'In Progress'.",
        "Resolve Ticket: Select an 'In Progress' complaint, click 'Upload Resolution Proof', select a sample/file photo of the repaired work, type resolution notes, and click 'Submit Resolution Evidence'."
    ]
    for s in o_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(s).font.size = Pt(10)

    # Workflow 3
    doc.add_heading("3.3 Citizen Verification Loop", level=2)
    v_steps = [
        "Notification Alert: The citizen receives a notification when their issue is marked 'Resolved'.",
        "Inspect Proof: Citizen opens the ticket modal and reviews the department's After-Photo proof.",
        "Approve / Close: If satisfied, click 'Confirm & Close Ticket' (+20 bonus points). Status transitions to 'Closed'.",
        "Reject / Reopen: If the fix is inadequate, click 'Reject & Reopen'. Status reverts to 'Reopened' with automatic alert to management."
    ]
    for s in v_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(s).font.size = Pt(10)

    # Workflow 4
    doc.add_heading("3.4 Executive Command Center (Admin)", level=2)
    a_steps = [
        "Sign In: Log in as 'Admin Center' (admin@civicflow.org).",
        "Command Center KPIs: Review Total System Complaints, Citizen Signups, Resolution Rate %, Overdue SLA Files, and Breach Escalations.",
        "Visual Analytics: Inspect Volume Distribution Bar Charts and Lifecycle Stage Pie Charts.",
        "Predictive Insights: View AI Spot Alerts and recurrent infrastructure warnings.",
        "Case Inspection: Scroll down to 'Breached SLA Escalation Logs' and click 'Inspect Case' to review any escalated file."
    ]
    for s in a_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(s).font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # SLA Table
    h4 = doc.add_heading("4. SLA Severity & Escalation Matrix", level=1)
    h4.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    sla_table = doc.add_table(rows=5, cols=4)
    sla_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sla_headers = ["Severity Level", "Target SLA Deadline", "Escalation Level 1", "Escalation Level 2"]
    for idx, ht in enumerate(sla_headers):
        cell = sla_table.rows[0].cells[idx]
        cell.text = ht
        set_cell_background(cell, "0F172A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)

    sla_rows = [
        ("Critical", "24 Hours (1 Day)", "Supervisor (Immediate Breach)", "Higher Authority / Commissioner (>48h)"),
        ("High", "72 Hours (3 Days)", "Supervisor (+1h post breach)", "Higher Authority (+48h)"),
        ("Medium", "168 Hours (7 Days)", "Supervisor (+1h post breach)", "Higher Authority (+48h)"),
        ("Low", "336 Hours (14 Days)", "Supervisor (+1h post breach)", "Higher Authority (+48h)")
    ]

    for r_idx, r_data in enumerate(sla_rows, start=1):
        cells = sla_table.rows[r_idx].cells
        for c_idx, val in enumerate(r_data):
            cells[c_idx].text = val
            if r_idx % 2 == 1:
                set_cell_background(cells[c_idx], "F1F5F9")
            for p in cells[c_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    doc.save(DOCX_PATH)
    print("Created DOCX:", DOCX_PATH)


def create_pdf_guide():
    doc = SimpleDocTemplate(PDF_PATH, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#1E3A8A'),
        alignment=1,
        spaceAfter=6
    )
    sub_style = ParagraphStyle(
        'DocSub',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#475569'),
        alignment=1,
        spaceAfter=20
    )
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=14,
        spaceAfter=8
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=10,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#334155'),
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    # Title & Subtitle
    story.append(Paragraph("CivicFlow Platform", title_style))
    story.append(Paragraph("Official User Guide & Operational Manual", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#1E3A8A'), spaceAfter=15))

    # Section 1
    story.append(Paragraph("1. Executive Overview", h1_style))
    story.append(Paragraph(
        "CivicFlow is a state-of-the-art civic accountability platform turning public infrastructure reports into trackable, "
        "measurable, and verified community action. Featuring AI-assisted classification, interactive Leaflet GIS mapping, "
        "dynamic SLA deadline countdowns, automatic supervisor escalations, and mandatory citizen resolution verification.",
        body_style
    ))

    # Section 2: Roles Table
    story.append(Paragraph("2. Access Roles & Demo Credentials", h1_style))
    table_data = [
        [Paragraph("<b>Role Type</b>", body_style), Paragraph("<b>Demo Email</b>", body_style), Paragraph("<b>Password</b>", body_style), Paragraph("<b>Responsibilities</b>", body_style)],
        [Paragraph("Citizen Reporter", body_style), Paragraph("citizen@civicflow.org", body_style), Paragraph("password", body_style), Paragraph("Report issues, upvote items, verify fixes.", body_style)],
        [Paragraph("Department Officer", body_style), Paragraph("authority@civicflow.org", body_style), Paragraph("password", body_style), Paragraph("Accept tickets, upload proof photos, resolve.", body_style)],
        [Paragraph("Platform Admin", body_style), Paragraph("admin@civicflow.org", body_style), Paragraph("password", body_style), Paragraph("Command Center KPIs, SLA escalations audit.", body_style)]
    ]
    t = Table(table_data, colWidths=[110, 130, 75, 195])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E3A8A')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('BACKGROUND', (0,1), (-1,1), colors.HexColor('#F8FAFC')),
        ('BACKGROUND', (0,3), (-1,3), colors.HexColor('#F8FAFC')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    # Section 3: Workflows
    story.append(Paragraph("3. User Workflows & Instructions", h1_style))
    
    story.append(Paragraph("3.1 How Citizens Report Issues", h2_style))
    c_steps = [
        "<b>1. Sign In:</b> Click <i>Sign In</i> ➔ Choose <i>Citizen User</i> demo login.",
        "<b>2. Upload Evidence:</b> Click <i>Report Issue</i>. Attach photo via File Upload, Image URL, or click Sample Evidence Photo presets (Pothole, Garbage, Streetlight, etc.).",
        "<b>3. Pin Location:</b> Select exact location on interactive Leaflet GIS map.",
        "<b>4. AI Analysis:</b> Enter description and run AI classification (auto-tags category, severity & SLA).",
        "<b>5. Duplicate Check & Submit:</b> Upvote existing nearby issues or click Submit to generate tracking code (CF-YYYYMMDD-XXXX)."
    ]
    for cs in c_steps:
        story.append(Paragraph(cs, bullet_style))

    story.append(Spacer(1, 6))
    story.append(Paragraph("3.2 How Officers Resolve Issues", h2_style))
    o_steps = [
        "<b>1. Sign In:</b> Log in as <i>Officer User</i> (authority@civicflow.org).",
        "<b>2. Accept Ticket:</b> Open Dashboard, select an Assigned issue, and click <i>Accept Assignment</i> (status ➔ In Progress).",
        "<b>3. Resolution Proof:</b> Select issue, click <i>Upload Resolution Proof</i>, pick After-Photo proof, type notes, and submit."
    ]
    for os_step in o_steps:
        story.append(Paragraph(os_step, bullet_style))

    story.append(Spacer(1, 6))
    story.append(Paragraph("3.3 Citizen Verification & Escalations", h2_style))
    v_steps = [
        "<b>Verification:</b> Reporter inspects department's After-Photo. Approving closes ticket (+20 points). Rejecting reopens ticket.",
        "<b>SLA Escalations:</b> Overdue complaints automatically escalate to Department Supervisor (Level 1) and Municipal Commissioner (Level 2)."
    ]
    for vs in v_steps:
        story.append(Paragraph(vs, bullet_style))

    story.append(Spacer(1, 10))

    # Section 4: SLA Matrix
    story.append(Paragraph("4. SLA Target Deadlines", h1_style))
    sla_data = [
        [Paragraph("<b>Severity</b>", body_style), Paragraph("<b>SLA Deadline</b>", body_style), Paragraph("<b>Level 1 Escalation</b>", body_style), Paragraph("<b>Level 2 Escalation</b>", body_style)],
        [Paragraph("Critical", body_style), Paragraph("24 Hours", body_style), Paragraph("Supervisor (Immediate)", body_style), Paragraph("Higher Authority (>48h)", body_style)],
        [Paragraph("High", body_style), Paragraph("72 Hours", body_style), Paragraph("Supervisor (+1h post breach)", body_style), Paragraph("Higher Authority (+48h)", body_style)],
        [Paragraph("Medium", body_style), Paragraph("168 Hours (7 Days)", body_style), Paragraph("Supervisor (+1h post breach)", body_style), Paragraph("Higher Authority (+48h)", body_style)],
        [Paragraph("Low", body_style), Paragraph("336 Hours (14 Days)", body_style), Paragraph("Supervisor (+1h post breach)", body_style), Paragraph("Higher Authority (+48h)", body_style)]
    ]
    st = Table(sla_data, colWidths=[80, 110, 160, 160])
    st.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F172A')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('BACKGROUND', (0,1), (-1,1), colors.HexColor('#F1F5F9')),
        ('BACKGROUND', (0,3), (-1,3), colors.HexColor('#F1F5F9')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('TOPPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(st)

    doc.build(story)
    print("Created PDF:", PDF_PATH)

if __name__ == "__main__":
    create_docx_guide()
    create_pdf_guide()
